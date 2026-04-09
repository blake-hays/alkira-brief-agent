"""
Alkira Opportunity Brief Generator — Web App

Streamlit frontend for the Managed Agent. Partners type a company name,
the agent researches it and returns a formatted brief on the page.

Usage:
    streamlit run app.py
"""

import io
import os
import time
from dataclasses import dataclass

import streamlit as st
from anthropic import Anthropic
from dotenv import load_dotenv

load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))


# ── Configuration ────────────────────────────────────────────────

@dataclass(frozen=True)
class AgentConfig:
    agent_id: str
    env_id: str
    api_key: str


def _secret(key: str) -> str:
    """Read from env vars first, then st.secrets (for Streamlit Community Cloud)."""
    val = os.environ.get(key, "")
    if not val:
        try:
            val = st.secrets.get(key, "")
        except FileNotFoundError:
            pass
    return val


def load_config() -> AgentConfig:
    return AgentConfig(
        agent_id=_secret("ALKIRA_AGENT_ID"),
        env_id=_secret("ALKIRA_ENV_ID"),
        api_key=_secret("ANTHROPIC_API_KEY"),
    )


# ── Docx conversion ────────────────────────────────────────────

def _brief_to_docx(markdown_text: str) -> bytes:
    """Convert brief markdown to a basic .docx for download."""
    from docx import Document

    doc = Document()
    for line in markdown_text.split("\n"):
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped.startswith("|") and not all(c in "|-: " for c in stripped):
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                doc.add_paragraph("  |  ".join(cells))
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Agent Session ────────────────────────────────────────────────

def run_agent_session(
    config: AgentConfig,
    company_name: str,
    status_callback,
    timeout_seconds: float = 300,
) -> str:
    """
    Create a managed agent session, send the company name,
    stream events, and return the brief as markdown text.
    """
    client = Anthropic(api_key=config.api_key)

    prompt = (
        f'Research the company "{company_name}" and generate a one-page Alkira '
        f"opportunity brief. Follow the brief structure and writing style in your "
        f"instructions exactly. Focus on information from the past 12 months. "
        f"Return the brief as markdown text."
    )

    # Start session
    status_callback("Starting agent session...")
    session = client.beta.sessions.create(
        agent=config.agent_id,
        environment_id=config.env_id,
        title=f"Alkira Brief: {company_name}",
    )

    # Stream agent work
    brief_parts: list[str] = []
    status_callback(f"Researching {company_name}...")

    session_start = time.monotonic()

    with client.beta.sessions.events.stream(session.id) as stream:
        # Send user message after opening the stream (per docs)
        client.beta.sessions.events.send(
            session.id,
            events=[
                {
                    "type": "user.message",
                    "content": [{"type": "text", "text": prompt}],
                },
            ],
        )

        for event in stream:
            if time.monotonic() - session_start > timeout_seconds:
                raise TimeoutError(
                    f"Agent did not finish within {timeout_seconds / 60:.0f} minutes"
                )

            if not hasattr(event, "type"):
                continue

            # Capture text output
            if event.type == "agent.message":
                for block in getattr(event, "content", []):
                    if getattr(block, "type", None) == "text":
                        brief_parts.append(block.text)

            # Surface tool use as status updates
            if event.type == "agent.tool_use":
                tool_name = getattr(event, "name", "")
                if "search" in tool_name:
                    status_callback(f"Searching the web for {company_name} intel...")

            # Agent is done
            if event.type == "session.status_idle":
                break

            if event.type == "session.error":
                msg = getattr(getattr(event, "error", None), "message", "unknown")
                brief_parts.append(f"\n\n[Error: {msg}]")
                break

    return "".join(brief_parts)


# ── Streamlit UI ─────────────────────────────────────────────────

def main() -> None:
    st.set_page_config(
        page_title="Alkira Brief Generator",
        page_icon="🔷",
        layout="centered",
    )

    # ── Session state ────────────────────────────────────────
    if "brief_history" not in st.session_state:
        st.session_state.brief_history = []

    # ── Sidebar: recent briefs ───────────────────────────────
    with st.sidebar:
        st.markdown("### Recent Briefs")
        if not st.session_state.brief_history:
            st.caption("No briefs generated yet this session.")
        for i, entry in enumerate(st.session_state.brief_history):
            st.download_button(
                label=entry["company"],
                data=_brief_to_docx(entry["brief_md"]),
                file_name=entry["filename"],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"history_{i}",
            )

    # ── Custom CSS ───────────────────────────────────────────
    st.markdown("""
    <style>
        .main-header {
            font-size: 2rem;
            font-weight: 700;
            color: #1B4F72;
            margin-bottom: 0;
        }
        .sub-header {
            font-size: 1rem;
            color: #666;
            margin-top: 0;
            margin-bottom: 2rem;
        }
        .stButton > button {
            background-color: #1B4F72;
            color: white;
            border: none;
            padding: 0.6rem 2rem;
            font-size: 1rem;
            border-radius: 6px;
            width: 100%;
        }
        .stButton > button:hover {
            background-color: #2980B9;
        }
        .status-box {
            background-color: #EBF5FB;
            border-left: 4px solid #2980B9;
            padding: 1rem;
            border-radius: 4px;
            margin: 1rem 0;
        }
    </style>
    """, unsafe_allow_html=True)

    # ── Header ───────────────────────────────────────────────
    st.markdown('<p class="main-header">Alkira Opportunity Brief Generator</p>', unsafe_allow_html=True)
    st.markdown(
        '<p class="sub-header">Enter a company name. The agent researches it and builds a '
        'one-page Alkira opportunity brief you can hand to your customer.</p>',
        unsafe_allow_html=True,
    )

    # ── Config check ─────────────────────────────────────────
    config = load_config()
    if not config.agent_id or not config.env_id or not config.api_key:
        st.error(
            "Missing configuration. Make sure `.env` contains "
            "ANTHROPIC_API_KEY, ALKIRA_AGENT_ID, and ALKIRA_ENV_ID. "
            "Run `python setup_agent.py` first."
        )
        return

    # ── Input ────────────────────────────────────────────────
    col1, col2 = st.columns([3, 1])
    with col1:
        company_name = st.text_input(
            "Company name",
            placeholder="e.g. Mary Kay, Chevron, Walmart",
            label_visibility="collapsed",
        )
    with col2:
        generate = st.button("Generate Brief", use_container_width=True)

    # ── Generation ───────────────────────────────────────────
    if generate and company_name.strip():
        status_placeholder = st.empty()
        progress = st.progress(0)

        step = {"count": 0}

        def update_status(msg: str) -> None:
            step["count"] += 1
            pct = min(step["count"] * 15, 95)
            progress.progress(pct)
            status_placeholder.markdown(
                f'<div class="status-box">{msg}</div>',
                unsafe_allow_html=True,
            )

        start = time.time()

        try:
            with st.spinner(""):
                brief_md = run_agent_session(
                    config, company_name.strip(), update_status,
                )

            elapsed = time.time() - start
            progress.progress(100)
            status_placeholder.empty()

            st.success(f"Brief generated in {elapsed:.0f} seconds.")

            # ── Display the brief ────────────────────────────
            st.markdown(brief_md)

            # ── Download as .docx ────────────────────────────
            safe_name = company_name.strip().replace(" ", "_")
            st.download_button(
                label=f"Download {safe_name}_Alkira_Brief.docx",
                data=_brief_to_docx(brief_md),
                file_name=f"{safe_name}_Alkira_Brief.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

            # Save to history
            st.session_state.brief_history.insert(0, {
                "company": company_name.strip(),
                "brief_md": brief_md,
                "filename": f"{safe_name}_Alkira_Brief.docx",
            })
            st.session_state.brief_history = st.session_state.brief_history[:5]

        except TimeoutError:
            progress.empty()
            status_placeholder.empty()
            st.error("Generation timed out after 5 minutes. Please try again.")

        except Exception as exc:
            progress.empty()
            status_placeholder.empty()
            st.error(f"Something went wrong: {exc}")

    elif generate:
        st.warning("Enter a company name first.")

    # ── How it works ───────────────────────────��─────────────
    with st.expander("How it works"):
        st.markdown(
            "This tool uses an AI agent to research a company and generate a "
            "one-page Alkira opportunity brief.\n\n"
            "1. **Enter a company name** — any company you're pursuing or "
            "meeting with.\n"
            "2. **The agent researches it** — searching the web for recent "
            "news, IT leadership, cloud strategy, and network infrastructure.\n"
            "3. **You get a brief** — displayed on the page with a .docx "
            "download option. The brief maps what the agent found to "
            "Alkira's entry points, with proof points and a suggested next "
            "step.\n\n"
            "Briefs typically take 1-2 minutes to generate."
        )


if __name__ == "__main__":
    main()
